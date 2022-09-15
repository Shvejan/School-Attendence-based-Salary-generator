from django.shortcuts import render
from .models import employee
import json
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from pprint import pprint
import docx
from calendar import monthrange
import inflect
# Create your views here.
def home(request):
    data =  employee.objects.all()
    return render(request,'generate.html',{'data':data})


@csrf_exempt
def generate(request):
    if request.method=='POST':
            data = json.loads(request.body.decode("utf-8"))
            pprint(data)
            month = data['month']
            # working_days=data['working_days']
            attendence = data['json']
            doc = docx.Document("payslip new.docx")
            for k,v in attendence.items():

                obj = employee.objects.get(pk=k)

                create_payslip(doc,obj,month,v['leaves_taken'],v['Absent'])
                # print(len(doc.tables))
                # table1 = doc.tables[1]
                # table2 = doc.tables[2]
                # table3 = doc.tables[0]
                # table4 = doc.tables[3]
                # table5 = doc.tables[4]


                # table5.cell(0, 1).text = 'new value'
                #
                # op = k + ".doc"
                # doc.save(op)

    return JsonResponse({'foo':'bar'})





def create_payslip(doc,obj,month,leaves,absent):
    hra= str(round(obj.basic * 40/100))
    total_earnings=str(obj.basic+int(hra)+obj.ca+obj.ma+obj.sa)
    pf="-"
    proof_tax = '200'
    # proof_tax= total_earnings<="15000" ? "":total_earnings<="15000" ? "150":"200"
    income_tax= '500'
    total_deduction=''


    table1 = doc.tables[0]
    table2 = doc.tables[1]

    monthAndYear, days=month_name(month)

    table1.cell(0, 1).text = obj.name
    table1.cell(1, 1).text = obj.id
    table1.cell(2, 1).text = str(obj.join_date)
    table1.cell(3, 1).text = obj.designation
    table1.cell(4, 1).text = obj.department

    table1.cell(0, 3).text = obj.pan
    table1.cell(1, 3).text = obj.bank
    table1.cell(2, 3).text = str(obj.accNo)
    #table2.cell(3, 1).text = obj.pfNo
    table1.cell(3, 3).text = obj.esiNo
    table1.cell(4, 3).text = obj.uan
    #
    #
    table1.cell(0, 5).text = leaves
    table1.cell(1, 5).text = absent
    table1.cell(2, 5).text = str(obj.leaves_allowed)
    table1.cell(3, 5).text = str(days-int(leaves)-int(absent))
    table1.cell(4, 5).text = str(days)
    #
    table2.cell(2, 1).text = str(obj.basic)
    table2.cell(3, 1).text = hra
    table2.cell(4, 1).text = str(obj.sa)
    table2.cell(5, 1).text = str(obj.ca)
    table2.cell(6, 1).text = str(obj.ma)
    table2.cell(7, 1).text = '-'
    #table2.cell(8, 1).text = str(obj.basic+obj.sa+obj.ca+obj.ma+int(hra))
    #
    table2.cell(2, 2).text = str(round(obj.basic/days*(days-int(leaves)-int(absent))))
    table2.cell(3, 2).text = str(round(int(hra)/days*(days-int(leaves)-int(absent))))
    table2.cell(4, 2).text = str(round(obj.sa/days*(days-int(leaves)-int(absent))))
    table2.cell(5, 2).text = str(round(obj.ca/days*(days-int(leaves)-int(absent))))
    table2.cell(6, 2).text = str(round(obj.ma/days*(days-int(leaves)-int(absent))))
    table2.cell(7, 2).text = "-"
    table2.cell(8, 2).text = str(round(int(total_earnings)/days*(days-int(leaves)-int(absent))))
    #
    table2.cell(2, 4).text = pf
    table2.cell(3, 4).text = proof_tax
    table2.cell(4, 4).text = income_tax
    table2.cell(8, 4).text = str(int(proof_tax)+int(income_tax))
    converter = inflect.engine()

    doc.tables[1].cell(9,1).text ="                                                                          Net Pay for this month  :  "+str(int(table2.cell(8, 2).text)-int(table2.cell(8, 4).text))
    doc.tables[1].cell(10,1).text = "                                                                        "+converter.number_to_words(doc.tables[1].cell(9,1).text)


    for p in doc.paragraphs:
        if "%month%" in p.text:
            p.text = p.text.replace("%month%",monthAndYear)

    op = obj.name + ".doc"
    doc.save(op)



def month_name(date):
    month=int(date[5:])
    name=''
    days=monthrange(int(date[:4]), int(date[5:]))
    if(month==1):
        name= 'January '+date[:4]

    elif(month==2):
        name= 'February '+date[:4]

    elif(month==3):
        name= 'March '+date[:4]

    elif(month==4):
        name= 'April '+date[:4]

    elif(month==5):
        name= 'May '+date[:4]

    elif(month==6):
        name= 'June '+date[:4]

    elif(month==7):
        name= 'July '+date[:4]

    elif(month==8):
        name= 'August '+date[:4]

    elif(month==9):
        name= 'September '+date[:4]

    elif(month==10):
        name= 'October '+date[:4]

    elif(month==11):
        name= 'November '+date[:4]

    elif(month==12):
        name= 'December '+date[:4]




    return name,days[1]
