import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
table.alignment = WD_TABLE_ALIGNMENT.CENTER
doc = docx.Document("test1.docx")
#doc.tables[1].cell(0,0).text="test1"

doc.tables[1].cell(9,0).text="test1"
doc.tables[1].cell(10,0).text="test2"
doc.tables[1].cell(9,1).text="test3"
doc.tables[1].cell(10,1).text="test4"
doc.save("newtest.doc")
